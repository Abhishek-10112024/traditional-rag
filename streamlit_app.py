"""Streamlit UI for RAG application."""

# pyrefly: ignore [missing-import]
import streamlit as st
from typing import List, Dict, Tuple
import sys
from pathlib import Path
# pyrefly: ignore [missing-import]
from loguru import logger
import io
import re
import hashlib
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
import subprocess
import tempfile
# pyrefly: ignore [missing-import]
from pypdf import PdfReader
from docx import Document
# pyrefly: ignore [missing-import]
import pdfplumber

# Add project to path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from src.rag_pipeline import RAGPipeline
from config import settings, LOGS_DIR

# Configure logger
logger.add(f"{LOGS_DIR}/streamlit.log", rotation="500 MB")

# Page configuration
st.set_page_config(
    page_title="RAG Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main {
        padding: 2rem;
    }
    .stTabs [data-baseweb="tab-list"] button {
        font-size: 1rem;
    }
    .chat-message {
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 0.5rem 0;
    }
    .user-message {
        background-color: #e3f2fd;
        color: #1976d2;
    }
    .assistant-message {
        background-color: #f3e5f5;
        color: #7b1fa2;
    }
</style>
""", unsafe_allow_html=True)


class RAGApp:
    """Streamlit RAG Application."""
    
    def __init__(self):
        """Initialize the app."""
        self.rag_pipeline = None
        self.init_session_state()
    
    def init_session_state(self):
        """Initialize session state."""
        if "rag_pipeline" not in st.session_state:
            st.session_state.rag_pipeline = None
        
        if "chat_history" not in st.session_state:
            st.session_state.chat_history = []
        
        if "documents_loaded" not in st.session_state:
            st.session_state.documents_loaded = False
        
        if "doc_count" not in st.session_state:
            st.session_state.doc_count = 0

        if "startup_reset_done" not in st.session_state:
            st.session_state.startup_reset_done = False
    
    def initialize_pipeline(self):
        """Initialize RAG pipeline."""
        try:
            with st.spinner("Initializing RAG pipeline..."):
                self.rag_pipeline = RAGPipeline(
                    use_reranking=st.session_state.get("use_reranking", True),
                    use_hybrid_search=st.session_state.get("use_hybrid_search", True)
                )
                st.session_state.rag_pipeline = self.rag_pipeline
                st.success("✅ RAG pipeline initialized successfully!")
        except Exception as e:
            st.error(f"❌ Error initializing RAG pipeline: {e}")
            logger.error(f"Error initializing RAG pipeline: {e}")

    # ------------------------------------------------------------------
    # Smart chunking: table-aware + recursive text + heading prefix
    # ------------------------------------------------------------------

    def _is_heading_line(self, line: str) -> bool:
        """Return True if the line looks like a section heading."""
        stripped = line.strip()
        if not stripped or len(stripped) > 120:
            return False
        # Markdown headings (#, ##, …)
        if stripped.startswith('#'):
            return True
        # Bold-only line (**Title**)
        if re.match(r'^\*\*[^*]+\*\*\.?$', stripped):
            return True
        # ALL-CAPS short phrase (2–10 words, no punctuation at end)
        words = stripped.split()
        if stripped.isupper() and 2 <= len(words) <= 10:
            return True
        return False

    def _split_segments(self, text: str) -> list:
        """Split text into alternating TABLE and PROSE segments."""
        segments: list = []
        lines = text.split('\n')
        buf_type = 'prose'
        buf_lines: List[str] = []

        def flush(t: str, ls: List[str]):
            content = '\n'.join(ls).strip()
            if content:
                segments.append({'type': t, 'content': content})

        for line in lines:
            stripped = line.strip()
            # A line belongs to a table if it starts with '|'
            is_table_row = stripped.startswith('|')
            if is_table_row:
                if buf_type == 'prose':
                    flush('prose', buf_lines)
                    buf_lines = []
                buf_type = 'table'
                buf_lines.append(line)
            else:
                if buf_type == 'table':
                    flush('table', buf_lines)
                    buf_lines = []
                buf_type = 'prose'
                buf_lines.append(line)

        flush(buf_type, buf_lines)
        return segments

    def _chunk_prose(self, text: str, heading_prefix: str = '') -> List[str]:
        """Recursively chunk prose: paragraph → sentence → word fallback."""
        chunk_size = settings.CHUNK_SIZE
        overlap = min(settings.CHUNK_OVERLAP, chunk_size - 1)
        prefix = f"{heading_prefix} | " if heading_prefix else ''

        words = text.split()
        if not words:
            return []

        # Fits in a single chunk — return directly
        if len(words) <= chunk_size:
            return [f"{prefix}{text.strip()}"]

        # ── Level 1: paragraph split ──────────────────────────────────
        paragraphs = [p.strip() for p in re.split(r'\n{2,}', text) if p.strip()]
        if len(paragraphs) > 1:
            chunks: List[str] = []
            buffer: List[str] = []
            for para in paragraphs:
                para_words = para.split()
                if len(para_words) > chunk_size:
                    # Flush buffer first, then recurse on oversized paragraph
                    if buffer:
                        chunks.append(f"{prefix}{' '.join(buffer)}")
                        buffer = buffer[-overlap:] if overlap else []
                    sub = self._chunk_prose(para, heading_prefix)
                    chunks.extend(sub)
                    if sub:
                        buffer = sub[-1].split()[-overlap:] if overlap else []
                elif len(buffer) + len(para_words) <= chunk_size:
                    buffer.extend(para_words)
                else:
                    chunks.append(f"{prefix}{' '.join(buffer)}")
                    buffer = buffer[-overlap:] if overlap else []
                    buffer.extend(para_words)
            if buffer:
                chunks.append(f"{prefix}{' '.join(buffer)}")
            return [c for c in chunks if c.strip()]

        # ── Level 2: sentence split ───────────────────────────────────
        sentences = re.split(r'(?<=[.!?])\s+', text.strip())
        if len(sentences) > 1:
            chunks = []
            buffer = []
            for sent in sentences:
                sent_words = sent.split()
                if len(buffer) + len(sent_words) <= chunk_size:
                    buffer.extend(sent_words)
                else:
                    if buffer:
                        chunks.append(f"{prefix}{' '.join(buffer)}")
                        buffer = buffer[-overlap:] if overlap else []
                    buffer.extend(sent_words)
            if buffer:
                chunks.append(f"{prefix}{' '.join(buffer)}")
            return [c for c in chunks if c.strip()]

        # ── Level 3: word-level overlap (last resort) ─────────────────
        step = max(1, chunk_size - overlap)
        chunks = []
        for start in range(0, len(words), step):
            chunk = f"{prefix}{' '.join(words[start:start + chunk_size])}"
            if chunk.strip():
                chunks.append(chunk)
        return chunks

    def _smart_chunk(self, text: str) -> List[str]:
        """Table-aware smart chunker with section heading context.

        - Markdown table blocks → preserved as a single chunk each.
        - Prose → paragraph/sentence/word recursive split.
        - Section heading lines are detected and prepended to following chunks.
        """
        segments = self._split_segments(text)
        chunks: List[str] = []
        current_heading = ''

        for seg in segments:
            if seg['type'] == 'table':
                # Keep entire table as one chunk; prepend heading context
                table_text = seg['content']
                if current_heading:
                    table_text = f"{current_heading}\n{table_text}"
                chunks.append(table_text)
            else:
                # Prose: scan for headings, chunk blocks between headings
                prose_lines = seg['content'].split('\n')
                block_lines: List[str] = []
                block_heading = current_heading

                for line in prose_lines:
                    if self._is_heading_line(line):
                        # Flush accumulated prose before this heading
                        if block_lines:
                            block_text = '\n'.join(block_lines)
                            chunks.extend(self._chunk_prose(block_text, block_heading))
                            block_lines = []
                        # Update active heading
                        block_heading = line.strip().lstrip('#').strip().strip('*')
                        current_heading = block_heading
                    else:
                        block_lines.append(line)

                # Flush remaining prose
                if block_lines:
                    block_text = '\n'.join(block_lines)
                    chunks.extend(self._chunk_prose(block_text, block_heading))

        return [c for c in chunks if c.strip()]

    def _chunk_text(self, text: str) -> List[str]:
        """Legacy word-overlap chunker — delegates to _smart_chunk."""
        return self._smart_chunk(text)

    def _extract_pdf_text(self, file_bytes: bytes) -> str:
        """Fallback: flat text extraction via pypdf (no table structure)."""
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()

    def _extract_pdf_pages_structured(self, file_bytes: bytes, filename: str):
        """Page-aware PDF extraction using pdfplumber.

        For each page:
        - Tables are extracted as markdown (preserving rows/columns).
        - Remaining text is extracted from regions outside table areas.
        - Lone page-number lines (single digits) are filtered out.
        - A [Page N] marker is prepended so the LLM can cite sources.

        Returns:
            (texts, metadatas): parallel lists for RAGPipeline.add_documents().
            metadatas entries: {"source": filename, "page": int}
        """
        texts, metadatas = [], []
        try:
            # Build printed-page-label map via pypdf.
            # pypdf.page_labels gives the labels actually printed in the PDF
            # footer (e.g. 'i','ii','1','2') instead of the raw physical index.
            page_label_map: dict = {}
            try:
                _tmp_reader = PdfReader(io.BytesIO(file_bytes))
                for _idx, _lbl in enumerate(_tmp_reader.page_labels):
                    _lbl = (_lbl or "").strip()
                    page_label_map[_idx] = _lbl if _lbl else str(_idx + 1)
            except Exception as _e:
                logger.debug(f"page_labels unavailable ({_e}); using physical page index")

            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                # ── Pre-extract all tables via camelot (stream mode) ───────────
                # Stream mode detects columns by whitespace gaps — exactly how
                # government statistical reports format their data tables.
                # When camelot is available, we use it for tables and extract
                # ALL page text via pdfplumber (no outside_bbox filtering),
                # which avoids the text-loss problem.
                camelot_tables_by_page: dict = {}
                camelot_available = False
                _tmp_path = None
                try:
                    # pyrefly: ignore [missing-import]
                    import camelot as _camelot
                    import tempfile
                    import os as _os
                    with tempfile.NamedTemporaryFile(
                        suffix=".pdf", delete=False
                    ) as _tmp:
                        _tmp.write(file_bytes)
                        _tmp_path = _tmp.name
                    _all_camelot = _camelot.read_pdf(
                        _tmp_path,
                        pages="all",
                        flavor="stream",
                        edge_tol=50,
                        row_tol=10,
                    )
                    for _ct in _all_camelot:
                        camelot_tables_by_page.setdefault(_ct.page, []).append(_ct)
                    camelot_available = True
                    logger.debug(
                        f"Camelot found {len(_all_camelot)} tables in {filename}"
                    )
                except ImportError:
                    logger.debug("camelot not installed; using pdfplumber table detection")
                except Exception as _ce:
                    logger.warning(f"Camelot extraction failed ({_ce}); falling back to pdfplumber")
                finally:
                    if _tmp_path:
                        try:
                            _os.unlink(_tmp_path)
                        except Exception:
                            pass

                for page_num, page in enumerate(pdf.pages, start=1):
                    # Use printed label (e.g. '1', '52', 'i') when available
                    page_label = page_label_map.get(page_num - 1, str(page_num))
                    page_parts = [f"[Page {page_label}]"]

                    if camelot_available:
                        # ── Camelot path: reliable borderless table markdown ──────
                        for _ct in camelot_tables_by_page.get(page_num, []):
                            _df = _ct.df
                            # Skip trivial / malformed extractions
                            if _df is None or _df.empty:
                                continue
                            if _df.shape[0] < 2 or _df.shape[1] < 2:
                                continue
                            _rows = []
                            for _i, _row in _df.iterrows():
                                cells = [str(c).strip().replace("\n", " ") for c in _row]
                                _rows.append("| " + " | ".join(cells) + " |")
                                if _i == 0:
                                    _rows.append(
                                        "|" + "|".join(["---"] * len(cells)) + "|"
                                    )
                            page_parts.append("\n".join(_rows))

                        # Extract ALL page text (no outside_bbox filtering needed)
                        raw_text = (
                            page.extract_text(x_tolerance=3, y_tolerance=3) or ""
                        )

                    else:
                        # ── pdfplumber fallback (original behaviour) ──────────
                        table_bboxes = []
                        for table_obj in page.find_tables():
                            table_bboxes.append(table_obj.bbox)
                            data = table_obj.extract()
                            if not data:
                                continue
                            md_rows = []
                            for row_idx, row in enumerate(data):
                                cells = [
                                    str(c or "").strip().replace("\n", " ")
                                    for c in row
                                ]
                                md_rows.append("| " + " | ".join(cells) + " |")
                                if row_idx == 0:
                                    md_rows.append(
                                        "|" + "|".join(["---"] * len(cells)) + "|"
                                    )
                            page_parts.append("\n".join(md_rows))

                        filtered_page = page
                        for bbox in table_bboxes:
                            try:
                                filtered_page = filtered_page.outside_bbox(bbox)
                            except Exception:
                                pass
                        raw_text = (
                            filtered_page.extract_text(x_tolerance=3, y_tolerance=3)
                            or ""
                        )

                    # Filter lone page-number lines (e.g. "47", "  3  ")
                    clean_lines = [
                        line
                        for line in raw_text.split("\n")
                        if not line.strip().isdigit()
                    ]
                    clean_text = "\n".join(clean_lines).strip()
                    if clean_text:
                        page_parts.append(clean_text)

                    # 3. Build combined page content ------------------------
                    page_content = "\n\n".join(
                        p for p in page_parts if p.strip()
                    )
                    if page_content.strip() == f"[Page {page_label}]":
                        continue  # completely empty page

                    # 4. Chunk and tag with metadata -----------------------
                    for chunk in self._smart_chunk(page_content):
                        texts.append(chunk)
                        metadatas.append({"source": filename, "page": page_label})

        except Exception as e:
            logger.warning(
                f"pdfplumber failed for {filename}: {e} — falling back to pypdf"
            )
            raw = self._extract_pdf_text(file_bytes)
            for chunk in self._smart_chunk(raw):
                texts.append(chunk)
                metadatas.append({"source": filename, "page": 0})

        return texts, metadatas

    # LLaVA image extraction tunables
    _LLAVA_MIN_PX     = 200    # skip images smaller than this in either dimension
    _LLAVA_MAX_IMAGES = 20     # hard cap per document to avoid runaway processing
    _LLAVA_TIMEOUT    = 30     # seconds before giving up on one image

    def _extract_pdf_images(self, file_bytes: bytes, filename: str):
        """Extract and describe embedded images/charts from a PDF using pymupdf + LLaVA.

        Optimisations vs the original:
        - Min-size filter raised to 200×200 px (cuts icons/logos).
        - Hard cap of 20 figures per document.
        - MD5 deduplication — same image appearing on multiple pages is only
          described once.
        - 30-second per-image timeout via ThreadPoolExecutor; slow/hung images
          are skipped with a warning instead of blocking forever.
        - Streamlit progress bar so the user can see how far along we are.

        Requires: `ollama pull llava` to be run once before first use.
        """
        from src.llm import describe_image_bytes
        texts: List[str] = []
        metadatas: List[Dict] = []

        try:
            # pyrefly: ignore [missing-import]
            import fitz  # pymupdf
        except ImportError:
            logger.warning("pymupdf not installed; skipping image extraction.")
            return texts, metadatas

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")

            # ── Pass 1: collect candidate images ──────────────────────────
            candidates = []  # list of (page_num, img_idx, xref, img_bytes, w, h)
            seen_hashes: set = set()

            for page_num, page in enumerate(doc, start=1):
                for img_idx, img_info in enumerate(page.get_images(full=True), start=1):
                    if len(candidates) >= self._LLAVA_MAX_IMAGES:
                        break
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        width  = base_image.get("width", 0)
                        height = base_image.get("height", 0)
                        img_bytes = base_image["image"]

                        # Skip small images (icons, bullets, decorative)
                        if width < self._LLAVA_MIN_PX or height < self._LLAVA_MIN_PX:
                            continue

                        # Deduplicate: same content on multiple pages → describe once
                        img_hash = hashlib.md5(img_bytes).hexdigest()
                        if img_hash in seen_hashes:
                            logger.debug(
                                f"Skipped duplicate image xref={xref} p{page_num}"
                            )
                            continue
                        seen_hashes.add(img_hash)

                        candidates.append((page_num, img_idx, xref, img_bytes, width, height))
                    except Exception as e:
                        logger.debug(f"Skipped image xref={xref} p{page_num}: {e}")

                if len(candidates) >= self._LLAVA_MAX_IMAGES:
                    logger.info(
                        f"Reached {self._LLAVA_MAX_IMAGES}-image cap for {filename}; "
                        "remaining images skipped."
                    )
                    break

            if not candidates:
                return texts, metadatas

            # ── Pass 2: describe with timeout + progress bar ──────────────
            progress_bar = st.progress(0, text="🖼️ Describing figures with LLaVA…")

            for i, (page_num, img_idx, xref, img_bytes, width, height) in enumerate(candidates):
                progress_pct = int((i / len(candidates)) * 100)
                progress_bar.progress(
                    progress_pct,
                    text=f"🖼️ Figure {i + 1} of {len(candidates)} (page {page_num})…"
                )

                try:
                    with ThreadPoolExecutor(max_workers=1) as executor:
                        future = executor.submit(describe_image_bytes, img_bytes)
                        description = future.result(timeout=self._LLAVA_TIMEOUT)

                    if description:
                        texts.append(
                            f"[Page {page_num}] [Figure {img_idx}]\n{description}"
                        )
                        metadatas.append({
                            "source": filename,
                            "page": page_num,
                            "type": "figure",
                        })
                        logger.debug(
                            f"Described figure {img_idx} on page {page_num} "
                            f"({width}x{height}px)"
                        )
                except FuturesTimeoutError:
                    logger.warning(
                        f"LLaVA timeout ({self._LLAVA_TIMEOUT}s) for image on "
                        f"page {page_num} — skipped."
                    )
                except Exception as e:
                    logger.debug(f"Error describing image xref={xref} p{page_num}: {e}")

            progress_bar.progress(100, text=f"✅ Described {len(texts)} figure(s).")

        except Exception as e:
            logger.warning(f"Image extraction failed for {filename}: {e}")

        return texts, metadatas

    def _extract_docx_text(self, file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()

    def _extract_doc_text(self, file_bytes: bytes) -> str:
        """Extract legacy .doc text via system converters."""
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as tmp:
            tmp.write(file_bytes)
            tmp.flush()

            for command in (["antiword", tmp.name], ["catdoc", tmp.name]):
                try:
                    result = subprocess.run(
                        command,
                        capture_output=True,
                        text=True,
                        check=True
                    )
                    extracted = result.stdout.strip()
                    if extracted:
                        return extracted
                except (subprocess.CalledProcessError, FileNotFoundError):
                    continue

        raise ValueError(
            "Could not extract .doc text. Install `antiword` or `catdoc`, "
            "or convert the file to .docx."
        )

    def _extract_supported_text(self, file) -> str:
        file_bytes = file.getvalue()
        suffix = Path(file.name).suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf_text(file_bytes)
        if suffix == ".docx":
            return self._extract_docx_text(file_bytes)
        if suffix == ".doc":
            return self._extract_doc_text(file_bytes)

        raise ValueError(f"Unsupported file type: {suffix}")

    def _ensure_fresh_startup(self):
        """Reset persisted and in-memory indexes once per app session."""
        if st.session_state.startup_reset_done:
            return

        if self.rag_pipeline is None:
            self.initialize_pipeline()

        if self.rag_pipeline is not None:
            self.rag_pipeline.reset_knowledge_base(delete_storage=True)
            st.session_state.documents_loaded = False
            st.session_state.doc_count = 0
            st.session_state.chat_history = []
            st.session_state.startup_reset_done = True
            st.info("🧹 Started with a fresh knowledge base. Previous session data was removed.")
    
    def handle_document_upload(self):
        """Handle document upload."""
        st.subheader("📄 Document Management")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Upload Documents**")
            uploaded_files = st.file_uploader(
                "Upload research documents (.pdf, .doc, .docx)",
                type=["pdf", "doc", "docx"],
                accept_multiple_files=True
            )
            
            if uploaded_files:
                process_images = st.checkbox(
                    "🖼️ Describe charts & figures with LLaVA",
                    value=False,
                    help=(
                        "Uses the LLaVA vision model (via Ollama) to generate text "
                        "descriptions of charts and graphics found inside PDFs. "
                        "Run `ollama pull llava` once before enabling this. "
                        "Adds ~2-5 seconds per image during upload."
                    ),
                )
                if st.button("📤 Process and Index Documents", use_container_width=True):
                    with st.spinner("Processing documents..."):
                        try:
                            documents: List[str] = []
                            all_metadatas: List[Dict] = []

                            if self.rag_pipeline is None:
                                self.initialize_pipeline()
                            if self.rag_pipeline is None:
                                st.error("❌ Could not initialize RAG pipeline")
                                return

                            # Every new upload batch replaces prior indexed data.
                            self.rag_pipeline.reset_knowledge_base(delete_storage=True)
                            st.session_state.documents_loaded = False
                            st.session_state.doc_count = 0
                            st.session_state.chat_history = []

                            for file in uploaded_files:
                                file_bytes = file.getvalue()
                                suffix = Path(file.name).suffix.lower()
                                try:
                                    if suffix == ".pdf":
                                        # Structured extraction: tables + page tracking
                                        new_texts, new_metas = (
                                            self._extract_pdf_pages_structured(
                                                file_bytes, file.name
                                            )
                                        )
                                        if not new_texts:
                                            st.warning(
                                                f"⚠️ No extractable content in {file.name}"
                                            )
                                            continue
                                        documents.extend(new_texts)
                                        all_metadatas.extend(new_metas)

                                        # Optionally describe charts/figures with LLaVA
                                        if process_images:
                                            with st.spinner(
                                                f"🖼️ Describing figures in {file.name}..."
                                            ):
                                                img_texts, img_metas = (
                                                    self._extract_pdf_images(
                                                        file_bytes, file.name
                                                    )
                                                )
                                            if img_texts:
                                                documents.extend(img_texts)
                                                all_metadatas.extend(img_metas)
                                                st.info(
                                                    f"🖼️ Described {len(img_texts)} "
                                                    f"figure(s) from {file.name}"
                                                )
                                    elif suffix == ".docx":
                                        content = self._extract_docx_text(file_bytes)
                                        if not content.strip():
                                            st.warning(
                                                f"⚠️ No extractable text in {file.name}"
                                            )
                                            continue
                                        chunks = self._chunk_text(content)
                                        documents.extend(chunks)
                                        all_metadatas.extend(
                                            [{"source": file.name, "page": 0}]
                                            * len(chunks)
                                        )
                                    elif suffix == ".doc":
                                        content = self._extract_doc_text(file_bytes)
                                        if not content.strip():
                                            st.warning(
                                                f"⚠️ No extractable text in {file.name}"
                                            )
                                            continue
                                        chunks = self._chunk_text(content)
                                        documents.extend(chunks)
                                        all_metadatas.extend(
                                            [{"source": file.name, "page": 0}]
                                            * len(chunks)
                                        )
                                    else:
                                        st.warning(
                                            f"⚠️ Unsupported file type: {suffix}"
                                        )
                                        continue
                                except Exception as extract_error:
                                    st.warning(
                                        f"⚠️ Could not process {file.name}: {extract_error}"
                                    )
                                    logger.warning(f"Skipped {file.name}: {extract_error}")
                                    continue

                            if not documents:
                                st.error("❌ No documents were successfully processed")
                                return

                            self.rag_pipeline.add_documents(documents, all_metadatas)
                            st.session_state.documents_loaded = True
                            st.session_state.doc_count = len(documents)

                            st.success(
                                f"✅ Indexed {len(documents)} chunks from {len(uploaded_files)} file(s). "
                                "Previous indexed data was replaced."
                            )
                            logger.info(
                                f"Indexed {len(documents)} chunks via Streamlit (fresh replace mode)"
                            )
                            # Force a full re-render so the Chat tab unlocks
                            # immediately without requiring a second click.
                            st.rerun()
                        except Exception as e:
                            st.error(f"❌ Error processing documents: {e}")
                            logger.error(f"Error processing documents: {e}")
        
        with col2:
            st.write("**Dataset Information**")
            if st.session_state.documents_loaded:
                st.metric("Documents Indexed", st.session_state.doc_count)
                if st.button("🗑️ Clear All Documents", use_container_width=True):
                    try:
                        if self.rag_pipeline:
                            self.rag_pipeline.reset_knowledge_base(delete_storage=True)
                        st.session_state.documents_loaded = False
                        st.session_state.doc_count = 0
                        st.session_state.chat_history = []
                        st.success("✅ Documents cleared and persistent index reset.")
                    except Exception as e:
                        st.error(f"Error clearing documents: {e}")
            else:
                st.info("No documents indexed yet. Upload PDF/Word files to get started.")
    
    def handle_settings(self):
        """Handle settings panel."""
        st.sidebar.header("⚙️ Settings")
        
        with st.sidebar.expander("Search Configuration"):
            st.session_state.use_hybrid_search = st.checkbox(
                "Enable Hybrid Search (BM25 + Vector)",
                value=settings.HYBRID_SEARCH_ENABLED,
                help="Use both sparse and dense retrieval"
            )
            
            st.session_state.vector_weight = st.slider(
                "Vector Search Weight",
                0.0, 1.0, settings.VECTOR_WEIGHT,
                help="Weight for vector search in hybrid search"
            )
            
            st.session_state.top_k_retrieval = st.slider(
                "Top-K Documents to Retrieve",
                1, 30, settings.TOP_K_RETRIEVAL,
                help="Number of documents to retrieve before reranking"
            )
        
        with st.sidebar.expander("Reranking Configuration"):
            st.session_state.use_reranking = st.checkbox(
                "Enable Reranking",
                value=settings.RERANKING_ENABLED,
                help="Use cross-encoder for reranking"
            )
            
            st.session_state.top_k_rerank = st.slider(
                "Top-K Documents After Reranking",
                1, 15, settings.TOP_K_RERANK,
                help="Number of documents to pass to LLM"
            )
        
        with st.sidebar.expander("LLM Configuration"):
            st.session_state.temperature = st.slider(
                "Temperature",
                0.0, 2.0, settings.TEMPERATURE,
                step=0.1,
                help="Higher = more creative, Lower = more focused"
            )
            
            st.session_state.top_p = st.slider(
                "Top-P (Nucleus Sampling)",
                0.0, 1.0, settings.TOP_P,
                step=0.05,
                help="Controls diversity of output"
            )
            
            st.session_state.max_tokens = st.slider(
                "Max Tokens",
                1, 1000, settings.MAX_NEW_TOKENS,
                step=50,
                help="Maximum length of generated response"
            )
        
        with st.sidebar.expander("Cache Configuration"):
            st.session_state.cache_enabled = st.checkbox(
                "Enable Response Caching",
                value=settings.CACHE_ENABLED,
                help="Cache responses to identical queries"
            )
            
            if st.button("Clear Cache", use_container_width=True):
                if self.rag_pipeline:
                    self.rag_pipeline.clear_cache()
                    st.success("✅ Cache cleared!")
    
    def handle_chat(self):
        """Handle chat interface."""
        st.subheader("💬 Chat with RAG Assistant")
        
        # Check if pipeline is initialized
        if not st.session_state.documents_loaded:
            st.warning("⚠️ Please upload and index documents first!")
            return
        
        # Display chat history
        chat_container = st.container()
        with chat_container:
            for i, msg in enumerate(st.session_state.chat_history):
                if msg["role"] == "user":
                    with st.chat_message("user", avatar="👤"):
                        st.markdown(msg["content"])
                else:
                    with st.chat_message("assistant", avatar="🤖"):
                        st.markdown(msg["content"])
        
        # Input area
        st.divider()
        col1, col2 = st.columns([0.85, 0.15])
        
        with col1:
            user_input = st.text_input(
                "Ask a question:",
                key="user_input_",
                placeholder="What would you like to know?"
            )
        
        with col2:
            send_button = st.button("Send", use_container_width=True)
        
        # Process query
        if send_button and user_input:
            # Add user message to history
            st.session_state.chat_history.append({
                "role": "user",
                "content": user_input
            })
            
            # Generate response
            try:
                with st.spinner("Thinking..."):
                    # Initialize pipeline if not already done
                    if self.rag_pipeline is None:
                        self.initialize_pipeline()
                    
                    # Check if documents are loaded
                    has_documents = st.session_state.get("documents_loaded", False)
                    
                    if not has_documents:
                        st.warning("⚠️ Please upload and index documents first!")
                        st.info("💡 Upload PDF/Word files in the 'Document Management' tab.")
                        st.session_state.chat_history.pop()  # Remove the user message we just added
                    else:
                        # Apply runtime UI settings before query.
                        self.rag_pipeline.cache_enabled = st.session_state.get("cache_enabled", settings.CACHE_ENABLED)
                        if self.rag_pipeline.hybrid_search:
                            vector_weight = st.session_state.get("vector_weight", settings.VECTOR_WEIGHT)
                            self.rag_pipeline.hybrid_search.set_weights(
                                vector_weight=vector_weight,
                                bm25_weight=max(0.0, 1.0 - vector_weight)
                            )

                        result = self.rag_pipeline.query(
                            user_input,
                            top_k_retrieval=st.session_state.get("top_k_retrieval", settings.TOP_K_RETRIEVAL),
                            top_k_rerank=st.session_state.get("top_k_rerank", settings.TOP_K_RERANK),
                            temperature=st.session_state.get("temperature", settings.TEMPERATURE),
                            top_p=st.session_state.get("top_p", settings.TOP_P),
                            max_tokens=st.session_state.get("max_tokens", settings.MAX_NEW_TOKENS),
                            use_cache=st.session_state.get("cache_enabled", settings.CACHE_ENABLED),
                            # Last 3 exchanges (6 msgs), excluding the current user msg
                            # that was just appended — so the LLM sees prior context only
                            chat_history=st.session_state.chat_history[:-1][-6:],
                        )
                    
                        with st.expander("📚 Retrieved Documents"):
                            for i, doc in enumerate(result["retrieved_docs"], 1):
                                score = doc.get("score", "")
                                meta = doc.get("metadata", {})
                                page = meta.get("page") if isinstance(meta, dict) else None
                                source = meta.get("source", "") if isinstance(meta, dict) else ""

                                score_display = (
                                    f"{score:.3f}" if isinstance(score, (int, float)) else str(score)
                                )
                                page_display = f" | 📄 Page {page}" if page else ""
                                source_display = f" | {source}" if source else ""

                                st.markdown(
                                    f"**Chunk {i}** "
                                    f"(Score: {score_display}{page_display}{source_display})"
                                )
                                preview = doc["text"][:300]
                                st.write(preview + ("..." if len(doc["text"]) > 300 else ""))
                    
                        # Add assistant response
                        st.session_state.chat_history.append({
                            "role": "assistant",
                            "content": result["answer"]
                        })

                        # Display response
                        with st.chat_message("assistant", avatar="🤖"):
                            st.markdown(result["answer"])
                        
                        logger.info(f"Query processed: {user_input[:50]}...")
                        st.rerun()
            
            except Exception as e:
                st.error(f"❌ Error processing query: {e}")
                logger.error(f"Error processing query: {e}")
    
    def run(self):
        """Run the Streamlit app."""
        # Header
        st.title("🤖 RAG Assistant")
        st.markdown("*A fresh-session Retrieval-Augmented Generation system for PDF/Word research documents*")
        
        # Auto-initialize pipeline on page load
        if self.rag_pipeline is None and st.session_state.rag_pipeline is None:
            self.initialize_pipeline()

        self._ensure_fresh_startup()
        
        # Main layout
        tab1, tab2, tab3 = st.tabs(["💬 Chat", "📄 Documents", "ℹ️ About"])
        
        # Load settings
        self.handle_settings()
        
        with tab1:
            self.handle_chat()
        
        with tab2:
            self.handle_document_upload()
        
        with tab3:
            st.markdown("""
            ## About RAG Assistant

            A fully **local, offline-capable** Retrieval-Augmented Generation system built for
            government reports, research papers, and structured documents.

            ---

            ### ✨ Features

            | Capability | Detail |
            |---|---|
            | **Hybrid Search** | BM25 (sparse) + dense vector retrieval with configurable weights |
            | **Cross-Encoder Reranking** | `ms-marco-MiniLM-L-6-v2` re-scores candidates for precision |
            | **Smart Chunking** | Table-aware: markdown tables kept whole; prose split by paragraph → sentence → word with section-heading prefix |
            | **Multimodal (LLaVA)** | Embedded charts & figures described by LLaVA via Ollama; deduped, capped at 20/doc, 30 s timeout |
            | **Multi-turn Memory** | Last 3 exchanges passed to LLM so follow-up questions resolve correctly |
            | **Page Citations** | Every chunk tagged `[Page N]`; LLM instructed to cite page numbers |
            | **Response Cache** | In-memory answer cache (TTL 30 min) avoids re-generating identical queries |
            | **Fresh-session Reset** | On startup and each new upload batch, stale vectors are removed automatically |

            ---

            ### 🛠 Tech Stack

            | Layer | Technology |
            |---|---|
            | Embeddings | `sentence-transformers` — `all-MiniLM-L6-v2` |
            | Vector store | Chroma (cosine, persistent) |
            | Sparse retrieval | `rank-bm25` (BM25Okapi) |
            | Reranker | CrossEncoder (`ms-marco-MiniLM-L-6-v2`) |
            | PDF extraction | `pdfplumber` (structured) + `pypdf` (fallback) |
            | Image extraction | `pymupdf` (fitz) + LLaVA via Ollama |
            | LLM | Ollama (local) · HuggingFace Inference · Transformers (fallback) |
            | Frontend | Streamlit |
            | Config | `pydantic-settings` + `.env` |
            | Logging | loguru |

            ---

            ### 🚀 Getting Started

            1. Start **Ollama**: `ollama serve` → `ollama pull mistral`
            2. *(Optional for charts)* `ollama pull llava`
            3. Go to **📄 Documents** tab → upload PDF / DOCX → click **Process and Index**
            4. Switch to **💬 Chat** tab and ask questions

            ---

            ### 💡 Tips

            - Enable **LLaVA figure description** only if your PDF has meaningful charts (adds ~20 s per unique figure)
            - Raise **Top-K Retrieval** for broad questions; lower it for targeted look-ups
            - Use **Conversation History** naturally — "And what about Maharashtra?" works as a follow-up
            - **Reranking off** + broad query mode auto-activates for summary/overview questions
            - Logs: `logs/app.log` and `logs/streamlit.log`

            ---
            Built with ❤️ using free and open-source tools · Fully local · No API keys required
            """)


def main():
    """Main entry point."""
    app = RAGApp()
    app.run()


if __name__ == "__main__":
    main()
